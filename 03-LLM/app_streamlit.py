# 该py文件封装的是 给用户使用的前端页面.
import streamlit as st
import requests
import time
import pandas as pd
from io import BytesIO
from model2pred_Qwen import predict_fun
from collections import Counter

# 尝试导入 docx
try:
    from docx import Document
except ImportError:
    Document = None

# ---------- 页面配置 ----------
st.set_page_config(
    page_title="微博情绪 · 拾光阁",
    page_icon="🌸",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------- 自定义CSS（文艺风，增加细节） ----------
st.markdown("""
<style>
    body {
        font-family: 'Georgia', '宋体', 'KaiTi', serif;
        background-color: #fcf8f5;
        color: #4a3f3a;
    }
    .main {
        background: rgba(255, 248, 245, 0.75);
        backdrop-filter: blur(12px);
        border-radius: 24px;
        padding: 2rem 2.5rem;
        box-shadow: 0 8px 32px rgba(180, 140, 120, 0.15);
        margin: 1rem 0.5rem;
    }
    h1 {
        font-family: 'Georgia', '华文楷体', serif;
        font-weight: 400;
        color: #8b6b5e;
        letter-spacing: 4px;
        border-bottom: 2px dashed #e6d5cc;
        padding-bottom: 0.4rem;
    }
    h1::before {
        content: "🌸 ";
        font-size: 1.8rem;
    }
    h1::after {
        content: " 🌸";
        font-size: 1.8rem;
    }
    .subtitle {
        font-size: 1.1rem;
        color: #9f8b82;
        font-style: italic;
        margin-top: -0.5rem;
        margin-bottom: 1.5rem;
        text-align: center;
    }
    .intro-box {
        background: #faf3ef;
        border-radius: 20px;
        padding: 1.2rem 1.8rem;
        margin: 1.2rem 0;
        border-left: 6px solid #c9b0a3;
        font-size: 0.98rem;
        line-height: 1.7;
        color: #5d4a3e;
    }
    .tag-grid {
        display: flex;
        flex-wrap: wrap;
        gap: 8px 12px;
        justify-content: center;
        margin: 0.5rem 0 1rem 0;
    }
    .tag-item {
        background: #fffcf9;
        border-radius: 30px;
        padding: 0.3rem 1rem;
        border: 1px solid #e6d5cc;
        font-size: 0.9rem;
        color: #5d4a3e;
        box-shadow: 0 2px 6px rgba(180,140,120,0.06);
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 1.5rem;
        background: transparent;
        border-bottom: 2px solid #e6d5cc;
    }
    .stTabs [data-baseweb="tab"] {
        font-family: 'Georgia', serif;
        font-size: 1.1rem;
        color: #9f8b82;
        padding: 0.6rem 1.2rem;
        border-radius: 30px 30px 0 0;
        transition: 0.3s;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: #fffaf7;
        color: #7a5f52;
        border-bottom: 3px solid #c9b0a3;
        font-weight: 600;
    }
    .stTextInput > div > div > input {
        border-radius: 30px;
        border: 1px solid #dccfc8;
        padding: 0.8rem 1.2rem;
        background: #fffcf9;
        font-family: 'Georgia', serif;
        font-size: 1rem;
    }
    .stTextInput > div > div > input:focus {
        border-color: #c9b0a3;
        box-shadow: 0 0 0 2px rgba(201, 176, 163, 0.3);
    }
    .stButton > button {
        border-radius: 40px;
        background: #e8d9d0;
        color: #5d4a3e;
        border: none;
        padding: 0.6rem 2rem;
        font-family: 'Georgia', serif;
        font-size: 1rem;
        font-weight: 400;
        transition: all 0.3s ease;
        box-shadow: 0 4px 10px rgba(180, 140, 120, 0.15);
    }
    .stButton > button:hover {
        background: #dccbc0;
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(180, 140, 120, 0.25);
    }
    .stFileUploader > div {
        border: 2px dashed #dccfc8;
        border-radius: 24px;
        background: #fffcf9;
        padding: 1.5rem;
    }
    .dataframe {
        border-radius: 16px;
        overflow: hidden;
        background: #fffcf9;
        box-shadow: 0 4px 12px rgba(180, 140, 120, 0.08);
    }
    .dataframe th {
        background: #f5ede8 !important;
        color: #5d4a3e;
        font-weight: 500;
        font-family: 'Georgia', serif;
    }
    .stProgress > div > div {
        background: #dccbc0 !important;
        border-radius: 20px;
    }
    .stProgress > div {
        background: #f0e5df !important;
        border-radius: 20px;
    }
    .stAlert {
        border-radius: 20px;
        border-left: 6px solid #c9b0a3;
        background: #fffaf7;
    }
    .stRadio > div {
        gap: 1rem;
    }
    .stRadio label {
        font-family: 'Georgia', serif;
        color: #5d4a3e;
    }
    .example-chip {
        display: inline-block;
        background: #f5ede8;
        border-radius: 20px;
        padding: 0.2rem 0.8rem;
        margin: 0.2rem 0.4rem;
        font-size: 0.85rem;
        color: #5d4a3e;
        cursor: pointer;
        border: 1px solid #e6d5cc;
        transition: 0.2s;
    }
    .example-chip:hover {
        background: #e8d9d0;
    }
    hr {
        border: none;
        height: 1px;
        background: linear-gradient(to right, transparent, #e6d5cc, transparent);
        margin: 1.8rem 0;
    }
    .stat-box {
        background: #fffcf9;
        border-radius: 16px;
        padding: 0.8rem 1rem;
        margin: 0.5rem 0;
        border: 1px solid #f0e5df;
    }
</style>
""", unsafe_allow_html=True)

# ---------- 初始化 session_state ----------
if 'history' not in st.session_state:
    st.session_state.history = []


def add_history(text, pred_class, source="单条"):
    st.session_state.history.append({
        "text": text,
        "pred": pred_class,
        "time": time.strftime("%Y-%m-%d %H:%M:%S"),
        "source": source
    })


# ---------- 情绪标签映射（带Emoji） ----------
EMOJI_MAP = {
    "happy": "😊", "sad": "😢", "angry": "😠",
    "fear": "😨", "surprise": "😲", "neutral": "😐"
}
LABEL_DESC = {
    "happy": "快乐", "sad": "悲伤", "angry": "愤怒",
    "fear": "恐惧", "surprise": "惊讶", "neutral": "中性"
}

# ---------- 页面标题 ----------
st.markdown('<h1 style="text-align:center;">微博情绪 · 拾光阁</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">—— 于喧闹中，洞见每一缕心绪 ——</p>', unsafe_allow_html=True)

# ---------- 介绍与标签预览 ----------
st.markdown("""
<div class="intro-box">
    🌱 <b>情绪分类</b> 基于大语言模型，自动识别微博文本中的情绪倾向。<br>
    支持 <b>六种情绪</b>：快乐、悲伤、愤怒、恐惧、惊讶、中性。 
    无论是日常吐槽、深夜感慨，还是突发新闻，都能快速解读。
</div>
""", unsafe_allow_html=True)

# 标签展示
cols = st.columns(6)
for idx, (label, emoji) in enumerate(EMOJI_MAP.items()):
    with cols[idx]:
        st.markdown(
            f"<div style='text-align:center; background:#faf3ef; border-radius:30px; padding:4px 0; font-size:0.9rem;'>{emoji} {LABEL_DESC[label]}</div>",
            unsafe_allow_html=True)

st.markdown("---")

# ---------- 侧边栏：历史记录 + 统计 ----------
with st.sidebar:
    st.markdown("## 📜 时光手账")
    if st.session_state.history:
        # 统计情绪分布
        df_hist = pd.DataFrame(st.session_state.history)
        if 'pred' in df_hist.columns:
            counts = Counter(df_hist['pred'])
            total = len(df_hist)
            st.markdown("**📊 情绪分布**")
            for label, cnt in counts.items():
                pct = cnt / total * 100
                bar = "█" * int(pct // 2) + "░" * (50 - int(pct // 2))
                st.markdown(f"{EMOJI_MAP.get(label, '')} {LABEL_DESC.get(label, label)}: {cnt} ({pct:.0f}%)")
                st.progress(pct / 100, text="")
            st.divider()

        # 显示历史记录表格
        df_show = df_hist.iloc[::-1]  # 最新在前
        st.dataframe(
            df_show[["time", "text", "pred", "source"]],
            use_container_width=True,
            hide_index=True,
            height=300
        )
        if st.button("🗑️ 清空手账"):
            st.session_state.history = []
            st.rerun()
    else:
        st.info("🌸 尚无记录，请开始你的情绪探索")
        st.markdown("""
        <div style="background:#faf3ef; border-radius:16px; padding:12px; margin-top:8px; font-size:0.9rem;">
        💡 试试输入 “今天天气真好” → <b>happy</b><br>
        💡 或者 “气死我了” → <b>angry</b>
        </div>
        """, unsafe_allow_html=True)

# ---------- 主界面：两个 Tab ----------
tab1, tab2 = st.tabs(["🍃 单条细品", "📚 批量采撷"])

# ----- Tab1: 单条预测 -----
with tab1:
    st.markdown("### 输入一条微博，轻触分类")

    # 示例快速填充（点击后填入输入框，需要借助 session_state 或回调）
    # 使用 st.button 来设置示例文本（通过 session_state）
    if 'example_text' not in st.session_state:
        st.session_state.example_text = ""

    col_ex1, col_ex2, col_ex3 = st.columns(3)
    examples = ["今天阳光真好，心情都变好了", "又加班到深夜，好累", "什么？他竟然夺冠了！"]
    with col_ex1:
        if st.button("🌞 快乐示例", key="ex1"):
            st.session_state.example_text = examples[0]
            st.rerun()
    with col_ex2:
        if st.button("🌧️ 悲伤示例", key="ex2"):
            st.session_state.example_text = examples[1]
            st.rerun()
    with col_ex3:
        if st.button("😲 惊讶示例", key="ex3"):
            st.session_state.example_text = examples[2]
            st.rerun()

    # 文本输入框，绑定 session_state
    text = st.text_input('', value=st.session_state.example_text, placeholder="此刻，你想知道什么情绪？",
                         label_visibility="collapsed")

    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
    with col_btn2:
        if st.button('🌸 即刻探寻', key='single', use_container_width=True):
            if not text.strip():
                st.warning("请留下一些文字吧 ☕️")
            else:
                start_time = time.time()
                try:
                    result = requests.post("http://127.0.0.1:5000/predict", json={"text": text})
                    pred = result.json()['pred_class']
                    add_history(text, pred, source="单条")
                    st.success(f"✨ 耗时 {(time.time() - start_time) * 1000:.2f} ms")
                    emoji = EMOJI_MAP.get(pred, "")
                    label_cn = LABEL_DESC.get(pred, pred)
                    st.markdown(f"### 🏷️ 类别：{emoji} **{label_cn}** ({pred})")
                    # 清空示例状态
                    st.session_state.example_text = ""
                except Exception as e:
                    st.error('出错了，请稍后重试 🍂')
                    print(e)

# ----- Tab2: 批量预测 -----
with tab2:
    st.markdown("### 上传文件或粘贴多条文本，批量体悟")

    # 格式提示卡片
    with st.expander("📖 查看文件格式要求", expanded=False):
        st.markdown("""
        - **TXT**：每行一条文本，自动忽略空行。
        - **CSV**：建议格式为 `编号,文本,标签`（标签可选），程序自动取第二列。
        - **Excel (.xlsx)**：默认取第二列作为文本，若只有一列则取该列。
        - **Word (.docx)**：自动提取所有段落文字（含表格）。
        """)
        st.code("""示例CSV内容（无标题行）：
0,今天天气真好，心情愉快,happy
1,气死我了，车还没来,angry
2,这部电影太感人了,sad
""", language="text")

    # 选择输入方式
    input_method = st.radio("选择你的方式", ("📁 上传文件", "✍️ 手动输入"), horizontal=True)

    texts = []
    source_label = ""

    if input_method == "📁 上传文件":
        uploaded_file = st.file_uploader(
            "支持格式：TXT、CSV、Excel (.xlsx)、Word (.docx)",
            type=['txt', 'csv', 'xlsx', 'docx'],
            help="TXT/CSV: 纯文本；Excel/Word: 自动提取文字"
        )
        if uploaded_file is not None:
            file_type = uploaded_file.name.split('.')[-1].lower()
            st.info(f"📄 检测到文件：{uploaded_file.name}")
            source_label = f"批量-文件({file_type})"

            try:
                if file_type == 'txt':
                    content = uploaded_file.getvalue().decode('utf-8')
                    texts = [line.strip() for line in content.splitlines() if line.strip()]

                elif file_type == 'csv':
                    content = uploaded_file.getvalue().decode('utf-8')
                    lines = content.splitlines()
                    if not lines:
                        st.warning("文件为空")
                    else:
                        first_line = lines[0]
                        if any(kw in first_line for kw in ['编号', '文本', '情绪', '标签', 'id', 'text', 'label']):
                            data_lines = lines[1:]
                        else:
                            data_lines = lines
                        for line in data_lines:
                            parts = line.strip().split(',')
                            if len(parts) >= 2:
                                text_cell = parts[1].strip()
                                if text_cell:
                                    texts.append(text_cell)
                        if not texts:
                            st.error("未提取到文本，请确认CSV格式为：编号,文本,标签")

                elif file_type == 'xlsx':
                    df_excel = pd.read_excel(uploaded_file, header=None)
                    if df_excel.shape[0] > 0:
                        first_row = df_excel.iloc[0].astype(str).tolist()
                        first_row_str = ','.join(first_row)
                        if any(kw in first_row_str for kw in ['编号', '文本', '情绪', '标签', 'id', 'text', 'label']):
                            df_excel = df_excel.iloc[1:]
                    if df_excel.shape[1] == 1:
                        texts = df_excel.iloc[:, 0].astype(str).tolist()
                    else:
                        texts = df_excel.iloc[:, 1].astype(str).tolist()
                    texts = [t for t in texts if t.strip()]

                elif file_type == 'docx':
                    if Document is None:
                        st.error("未安装 python-docx，请运行：pip install python-docx")
                    else:
                        doc = Document(BytesIO(uploaded_file.getvalue()))
                        texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                        if not texts and doc.tables:
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        cell_text = cell.text.strip()
                                        if cell_text:
                                            texts.append(cell_text)
                else:
                    st.error(f"不支持的文件格式：{file_type}")

            except Exception as e:
                st.error(f"读取文件时出错：{str(e)}")
                print(f"详细错误：{e}")

            if texts:
                st.success(f"✅ 成功读取 {len(texts)} 条文本")
                st.write("预览（前5条）：")
                for i, t in enumerate(texts[:5]):
                    st.write(f"**{i + 1}.** {t[:80]}{'…' if len(t) > 80 else ''}")
            else:
                st.warning("未能读取到任何文本，请检查文件内容")

    else:  # 手动输入
        manual_text = st.text_area(
            "每行一条，轻轻写下：",
            height=200,
            placeholder="例如：\n今天天气真好\n气死我了，车还没来\n这部电影太感人了"
        )
        if manual_text:
            texts = [line.strip() for line in manual_text.splitlines() if line.strip()]
            if texts:
                st.info(f"📝 共输入 {len(texts)} 条文本")
                st.write("预览（前5条）：")
                for i, t in enumerate(texts[:5]):
                    st.write(f"**{i + 1}.** {t[:80]}{'…' if len(t) > 80 else ''}")
            source_label = "批量-手动"

    # 批量预测按钮
    if texts:
        if not source_label:
            source_label = "批量-未知"
        col_btn_b1, col_btn_b2, col_btn_b3 = st.columns([1, 1, 1])
        with col_btn_b2:
            if st.button('🌟 开启批量探寻', key='batch', use_container_width=True):
                progress_bar = st.progress(0)
                status_text = st.empty()

                results = []
                total = len(texts)
                for i, txt in enumerate(texts):
                    status_text.text(f"正在品味第 {i + 1}/{total} 条……")
                    res = predict_fun({'text': txt})
                    pred = res['pred_class']
                    add_history(txt, pred, source=source_label)
                    results.append(res)
                    progress_bar.progress((i + 1) / total)

                status_text.text("✨ 全部完成！")
                result_df = pd.DataFrame({
                    '文本': texts,
                    '预测类别': [r['pred_class'] for r in results],
                    '耗时(ms)': [f"{r['cost_time']:.2f}" for r in results]
                })

                st.success("🎉 批量预测完成！")
                st.dataframe(result_df, use_container_width=True, hide_index=True)

                csv_download = result_df.to_csv(index=False, encoding='utf-8-sig')
                st.download_button(
                    label="📥 下载结果 CSV",
                    data=csv_download,
                    file_name=f"predict_results_{int(time.time())}.csv",
                    mime="text/csv",
                    use_container_width=True
                )